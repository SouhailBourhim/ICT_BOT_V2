"""
Parser de documents pour différents formats
Support: PDF, TXT, MD, DOCX
"""
from pathlib import Path
from typing import Dict, List, Optional
from dataclasses import dataclass
from datetime import datetime
import re

import pypdf
from docx import Document as DocxDocument
import markdown
from bs4 import BeautifulSoup
from loguru import logger


@dataclass
class ParsedDocument:
    """Structure pour un document parsé"""
    content: str
    metadata: Dict
    pages: Optional[List[Dict]] = None  # Pour PDFs avec structure page


class DocumentParser:
    """Parser universel pour différents formats de documents"""
    
    SUPPORTED_FORMATS = {'.pdf', '.txt', '.md', '.docx'}
    
    def __init__(self):
        self.parsers = {
            '.pdf': self._parse_pdf,
            '.txt': self._parse_txt,
            '.md': self._parse_markdown,
            '.docx': self._parse_docx,
        }
    
    def parse(self, file_path: Path) -> ParsedDocument:
        """
        Parse un document selon son format
        
        Args:
            file_path: Chemin vers le document
            
        Returns:
            ParsedDocument avec contenu et métadonnées
        """
        if not file_path.exists():
            raise FileNotFoundError(f"Fichier introuvable: {file_path}")
        
        suffix = file_path.suffix.lower()
        if suffix not in self.SUPPORTED_FORMATS:
            raise ValueError(f"Format non supporté: {suffix}")
        
        logger.info(f"Parsing du document: {file_path.name}")
        
        # Parser selon le format
        parser_func = self.parsers[suffix]
        content, metadata, pages = parser_func(file_path)
        
        # Métadonnées communes
        base_metadata = self._extract_base_metadata(file_path)
        metadata.update(base_metadata)
        
        # Nettoyage du contenu
        content = self._clean_text(content)
        
        logger.success(f"✅ Document parsé: {len(content)} caractères")
        
        return ParsedDocument(
            content=content,
            metadata=metadata,
            pages=pages
        )
    
    def _parse_pdf(self, file_path: Path) -> tuple:
        """Parse un PDF avec extraction page par page"""
        pages_data = []
        all_text = []
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)
                
                metadata = {
                    'format': 'pdf',
                    'num_pages': len(pdf_reader.pages),
                    'pdf_metadata': pdf_reader.metadata or {}
                }
                
                for i, page in enumerate(pdf_reader.pages, 1):
                    page_text = page.extract_text()
                    
                    pages_data.append({
                        'page_number': i,
                        'content': page_text,
                        'char_count': len(page_text)
                    })
                    
                    all_text.append(f"\n--- Page {i} ---\n{page_text}")
                
                content = '\n'.join(all_text)
                return content, metadata, pages_data
                
        except Exception as e:
            logger.error(f"Erreur lors du parsing PDF: {e}")
            raise
    
    def _parse_txt(self, file_path: Path) -> tuple:
        """Parse un fichier texte simple"""
        try:
            # Détection de l'encodage (UTF-8, latin-1)
            encodings = ['utf-8', 'latin-1', 'cp1252']
            content = None
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()
                    break
                except UnicodeDecodeError:
                    continue
            
            if content is None:
                raise ValueError("Impossible de décoder le fichier texte")
            
            metadata = {
                'format': 'txt',
                'encoding': encoding,
                'line_count': content.count('\n')
            }
            
            return content, metadata, None
            
        except Exception as e:
            logger.error(f"Erreur lors du parsing TXT: {e}")
            raise
    
    def _parse_markdown(self, file_path: Path) -> tuple:
        """Parse un fichier Markdown et extrait le texte"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                md_content = file.read()
            
            # Conversion MD -> HTML -> Texte
            html = markdown.markdown(md_content, extensions=['extra', 'codehilite'])
            soup = BeautifulSoup(html, 'html.parser')
            
            # Extraction de la structure (titres)
            headers = [h.get_text() for h in soup.find_all(['h1', 'h2', 'h3'])]
            
            # Texte complet
            content = soup.get_text(separator='\n')
            
            metadata = {
                'format': 'markdown',
                'headers': headers,
                'has_code_blocks': '```' in md_content
            }
            
            return content, metadata, None
            
        except Exception as e:
            logger.error(f"Erreur lors du parsing Markdown: {e}")
            raise
    
    def _parse_docx(self, file_path: Path) -> tuple:
        """Parse un document Word DOCX"""
        try:
            doc = DocxDocument(file_path)
            
            # Extraction des paragraphes
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            content = '\n\n'.join(paragraphs)
            
            # Extraction des tableaux si présents
            tables_count = len(doc.tables)
            
            metadata = {
                'format': 'docx',
                'paragraph_count': len(paragraphs),
                'tables_count': tables_count,
                'core_properties': {
                    'author': doc.core_properties.author or 'Unknown',
                    'title': doc.core_properties.title or '',
                    'subject': doc.core_properties.subject or '',
                }
            }
            
            return content, metadata, None
            
        except Exception as e:
            logger.error(f"Erreur lors du parsing DOCX: {e}")
            raise
    
    def _extract_base_metadata(self, file_path: Path) -> Dict:
        """Extrait les métadonnées de base du fichier"""
        stat = file_path.stat()
        
        return {
            'filename': file_path.name,
            'filepath': str(file_path.absolute()),
            'file_size': stat.st_size,
            'created_at': datetime.fromtimestamp(stat.st_ctime).isoformat(),
            'modified_at': datetime.fromtimestamp(stat.st_mtime).isoformat(),
            'extension': file_path.suffix.lower(),
        }
    
    def _clean_text(self, text: str) -> str:
        """Nettoie le texte extrait"""
        # Suppression des espaces multiples
        text = re.sub(r'\s+', ' ', text)
        
        # Suppression des lignes vides multiples
        text = re.sub(r'\n\s*\n', '\n\n', text)
        
        # Trim
        text = text.strip()
        
        return text
    
    def batch_parse(self, file_paths: List[Path]) -> List[ParsedDocument]:
        """Parse plusieurs documents"""
        results = []
        
        for file_path in file_paths:
            try:
                doc = self.parse(file_path)
                results.append(doc)
            except Exception as e:
                logger.error(f"Échec du parsing de {file_path.name}: {e}")
                continue
        
        logger.info(f"✅ {len(results)}/{len(file_paths)} documents parsés avec succès")
        return results


# Test du parser
if __name__ == "__main__":
    parser = DocumentParser()
    
    # Exemple d'utilisation
    # doc = parser.parse(Path("data/documents/cours_iot.pdf"))
    # print(f"Contenu: {doc.content[:500]}...")
    # print(f"Métadonnées: {doc.metadata}")